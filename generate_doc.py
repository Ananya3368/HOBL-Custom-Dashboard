"""Generate the HOBL Dashboard design document as a Word file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ───────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.color.rgb = RGBColor(0, 0x78, 0xD4)

SCREENSHOTS_DIR = r"C:\Users\t-anaagarwal\Pictures\Screenshots"
SCREENSHOT_LOGIN = os.path.join(SCREENSHOTS_DIR, "Screenshot 2026-06-14 172121.png")
SCREENSHOT_AUTH = os.path.join(SCREENSHOTS_DIR, "Screenshot 2026-06-14 172205.png")
SCREENSHOT_DASHBOARD = os.path.join(SCREENSHOTS_DIR, "Screenshot 2026-06-14 172221.png")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ── Title Page ───────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("HOBL Dashboard")
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0, 0x78, 0xD4)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Design Document")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x60, 0x5E, 0x5C)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Date: June 14, 2026\n").font.size = Pt(12)
meta.add_run("Version: 1.0\n").font.size = Pt(12)

doc.add_page_break()

# ── Table of Contents ────────────────────────────────────────────────────────
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Overview",
    "2. End-to-End Data Flow",
    "3. Architecture",
    "4. Authentication",
    "5. Dashboard UI",
    "6. Kusto Table Schema",
    "7. API Endpoints",
    "8. Project Structure",
    "9. How to Run",
    "10. Dependencies",
    "11. Future Enhancements",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. Overview ──────────────────────────────────────────────────────────────
doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "The HOBL Dashboard is a web-based tool that visualizes performance and power metrics "
    "collected from HOBL (Hardware-Optimized Benchmark Lab) scenario runs. It queries the "
    "Fungates Kusto database and presents the data in an interactive, filterable HTML table."
)
doc.add_paragraph("The dashboard allows engineers to:")
items = [
    "Select a device and its RAM configuration",
    "Choose a scenario (e.g., YouTube, Teams, etc.)",
    "Pick a date when the scenario was run",
    "View all metrics iteration-wise in a structured table",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

# ── 2. End-to-End Data Flow ──────────────────────────────────────────────────
doc.add_heading("2. End-to-End Data Flow", level=1)
doc.add_paragraph(
    "The dashboard is the final consumer in a pipeline that starts with HOBL running a scenario:"
)

add_code_block(doc,
    "HOBL runs scenario (PASS)\n"
    "        │\n"
    "        ▼\n"
    "Results extracted to JSON + ETL, uploaded via CSEUploader.exe\n"
    "        │\n"
    "        ▼\n"
    "Fungates Azure Function ingests metrics into Kusto DB (Hobl_RawMetrics)\n"
    "        │\n"
    "        ▼\n"
    "HOBL Dashboard queries Kusto & displays data in browser"
)

doc.add_heading("Step-by-step:", level=2)

doc.add_paragraph(
    'Step 1 – HOBL (C:\\HOBL) runs a scenario on a device. If the result is PASS and '
    'fungates_upload_enabled=1, it automatically:'
)
doc.add_paragraph(
    "Extracts results into hobl_result.json (via utilities\\extractor\\json_builder.py)",
    style="List Bullet",
)
doc.add_paragraph(
    "Uploads hobl_result.json + trace.etl to the Fungates portal "
    "(via Upload-ToFungates.ps1 → CSEUploader.exe)",
    style="List Bullet",
)

doc.add_paragraph(
    "Step 2 – Fungates Azure Function (fungates-hobl) picks up the upload:"
)
doc.add_paragraph("Downloads the JSON and ETL blobs", style="List Bullet")
doc.add_paragraph(
    "Runs HoblMetricExtractor.exe to parse hobl_result.json into CSV metrics",
    style="List Bullet",
)
doc.add_paragraph(
    "Writes the metrics to the Kusto table Hobl_RawMetrics in the FungatesDataStore database",
    style="List Bullet",
)

doc.add_paragraph(
    "Step 3 – HOBL Dashboard (C:\\hobl-dashboard) queries the Kusto table and renders "
    "the data in the browser."
)

# ── 3. Architecture ──────────────────────────────────────────────────────────
doc.add_heading("3. Architecture", level=1)

add_code_block(doc,
    "┌─────────────────────────────────────────┐\n"
    "│           User's Browser                │\n"
    "│                                         │\n"
    "│   HOBL Dashboard (HTML + CSS + JS)      │\n"
    "│   ┌──────────┐ ┌────────┐ ┌──────┐     │\n"
    "│   │Device/RAM│ │Scenario│ │ Date │     │\n"
    "│   └────┬─────┘ └───┬────┘ └──┬───┘     │\n"
    "│        └──── cascading ───────┘         │\n"
    "│                  │                      │\n"
    "│        ┌─────────▼──────────┐           │\n"
    "│        │  Metrics Table     │           │\n"
    "│        └────────────────────┘           │\n"
    "└────────────────┬────────────────────────┘\n"
    "                 │  HTTP (localhost:5000)\n"
    "      ┌──────────▼───────────┐\n"
    "      │   Flask Backend      │\n"
    "      │   (app.py)           │\n"
    "      └──────────┬───────────┘\n"
    "                 │  KQL (azure-kusto-data)\n"
    "      ┌──────────▼───────────┐\n"
    "      │  Azure Data Explorer │\n"
    "      │  Cluster: fungateprd │\n"
    "      │  DB: FungatesDataStore│\n"
    "      │  Table: Hobl_RawMetrics│\n"
    "      └──────────────────────┘"
)

doc.add_heading("Tech Stack", level=2)
add_table(doc,
    ["Component", "Technology"],
    [
        ["Backend", "Python 3 + Flask"],
        ["Kusto SDK", "azure-kusto-data (KQL queries)"],
        ["Authentication", "azure-identity (InteractiveBrowserCredential)"],
        ["Frontend", "HTML + CSS + vanilla JavaScript"],
        ["Data Source", "Azure Data Explorer (Kusto)"],
    ],
)

# ── 4. Authentication ────────────────────────────────────────────────────────
doc.add_heading("4. Authentication", level=1)
doc.add_paragraph(
    "The dashboard uses Azure AD interactive browser login to authenticate the user. "
    "This ensures that only authorized users with Kusto RBAC access can view the data."
)

doc.add_heading("How it works:", level=2)

doc.add_paragraph(
    "1. When the Flask app starts, it initializes an InteractiveBrowserCredential "
    "from the azure-identity SDK."
)
doc.add_paragraph(
    '2. On the first Kusto query, a browser tab opens automatically showing the '
    'Microsoft Azure "Pick an account" login page:'
)

if os.path.exists(SCREENSHOT_LOGIN):
    doc.add_picture(SCREENSHOT_LOGIN, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 1: Azure AD Login — Pick an account")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x60, 0x5E, 0x5C)

doc.add_paragraph(
    '3. After successful login, the browser shows "Authentication complete. '
    'You can close this window.":'
)

if os.path.exists(SCREENSHOT_AUTH):
    doc.add_picture(SCREENSHOT_AUTH, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 2: Authentication Complete confirmation")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x60, 0x5E, 0x5C)

doc.add_paragraph(
    "4. The token is cached for the session — subsequent queries do not require re-login."
)

doc.add_heading("Access Control:", level=2)
doc.add_paragraph(
    "Only users with Azure AD accounts that have RBAC access to the Kusto cluster "
    "(fungateprd.centralus) can query data.",
    style="List Bullet",
)
doc.add_paragraph(
    "The dashboard runs on localhost (127.0.0.1), so only the person running it "
    "on their machine can access the web UI.",
    style="List Bullet",
)
doc.add_paragraph("No data is exposed externally.", style="List Bullet")

# ── 5. Dashboard UI ──────────────────────────────────────────────────────────
doc.add_heading("5. Dashboard UI", level=1)

if os.path.exists(SCREENSHOT_DASHBOARD):
    doc.add_picture(SCREENSHOT_DASHBOARD, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(
        "Figure 3: HOBL Dashboard — filters and status area (no data uploaded yet)"
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x60, 0x5E, 0x5C)

doc.add_heading("5.1 Filters (Cascading Dropdowns)", level=2)
doc.add_paragraph(
    "The dashboard has three cascading filter dropdowns that load data progressively:"
)

add_table(doc,
    ["Filter", "Source Column", "Description"],
    [
        ["Device & RAM Config", "TierName", 'Format: <DeviceName>_<MemorySizeGB>. Displayed as "DeviceName (XX GB)"'],
        ["Scenario", "TestName", "Extracted from TestName by removing the _NNN iteration suffix"],
        ["Date", "RunDate", "The date when the scenario was run, formatted as YYYY-MM-DD"],
    ],
)

doc.add_heading("Cascading behavior:", level=3)
doc.add_paragraph(
    "Device & RAM Config loads on page open → lists all available devices",
    style="List Number",
)
doc.add_paragraph(
    "Scenario loads when a device is selected → shows scenarios run on that device",
    style="List Number",
)
doc.add_paragraph(
    "Date loads when a scenario is selected → shows dates that scenario was run",
    style="List Number",
)
doc.add_paragraph(
    "Metrics table loads when a date is selected → shows all metrics for that run",
    style="List Number",
)

doc.add_heading("5.2 Metrics Table", level=2)
doc.add_paragraph(
    "Once all three filters are selected, the metrics are displayed in a table:"
)

add_table(doc,
    ["Column", "Source Column", "Description"],
    [
        ["Iteration", "Iteration", "The iteration number (1, 2, 3...) — grouped with a section header"],
        ["Metric Name", "Name", "The metric identifier (e.g., pm_emi_cpu_cluster_0, system_power)"],
        ["Value", "Value", "The numeric metric value (e.g., 0.177, 4.68)"],
        ["Unit", "Unit", "Unit of measurement (e.g., W for Watts, ms for milliseconds)"],
        ["Type", "MetricType", "Color-coded: PowerMetrics (green), PowerCalculation (orange), PerfMetrics (blue)"],
    ],
)

doc.add_paragraph(
    "Metrics are grouped by iteration. A blue header row separates each iteration "
    '(e.g., "Iteration 1", "Iteration 2"). Within each iteration, metrics are sorted '
    "by type and then by name."
)

# ── 6. Kusto Table Schema ───────────────────────────────────────────────────
doc.add_heading("6. Kusto Table Schema", level=1)
doc.add_paragraph(
    "The dashboard reads from the Hobl_RawMetrics table in the FungatesDataStore database. "
    "Key columns used by the dashboard:"
)

add_table(doc,
    ["Column", "Type", "Used For"],
    [
        ["Name", "string", "Metric name displayed in the table"],
        ["MetricType", "string", "Metric category (PowerMetrics, PowerCalculation, PerfMetrics)"],
        ["StudyType", "string", 'Run type (e.g., "Power")'],
        ["Value", "real", "Metric value displayed in the table"],
        ["Unit", "string", "Unit of measurement"],
        ["TierName", "string", "<DeviceName>_<RAMSizeGB> — used for device filter"],
        ["TestName", "string", "<scenario>_<iteration> — used for scenario filter"],
        ["RunDate", "datetime", "Date of the run — used for date filter"],
        ["Iteration", "int", "Iteration number — used for table grouping"],
        ["TestResultId", "string", "Unique ID for the test result"],
    ],
)

doc.add_heading("How TierName maps to Device + RAM:", level=2)
doc.add_paragraph(
    "The TierName is generated during HOBL's upload process (New-MetadataFile.ps1) as:"
)
add_code_block(doc, 'TierDefinitionKey = "<DeviceName>_<MemorySizeGB>"')
doc.add_paragraph(
    "For example: ROM-MS-IDCLAB-2_16 → Device: ROM-MS-IDCLAB-2, RAM: 16 GB"
)

doc.add_heading("How TestName maps to Scenario + Iteration:", level=2)
doc.add_paragraph(
    "The TestName is the ETL filename without extension (e.g., youtube_001). "
    "The dashboard extracts:"
)
doc.add_paragraph("Scenario: everything before the last _NNN → youtube", style="List Bullet")
doc.add_paragraph("Iteration: the numeric suffix → 1", style="List Bullet")

# ── 7. API Endpoints ────────────────────────────────────────────────────────
doc.add_heading("7. API Endpoints", level=1)
doc.add_paragraph("The Flask backend exposes four REST endpoints:")

add_table(doc,
    ["Endpoint", "Method", "Parameters", "Returns"],
    [
        ["/api/tiers", "GET", "—", "List of distinct TierName values"],
        ["/api/scenarios", "GET", "tier (TierName)", "List of scenario names for that tier"],
        ["/api/dates", "GET", "tier, scenario", "List of run dates (YYYY-MM-DD)"],
        ["/api/metrics", "GET", "tier, scenario, date", "List of metric objects with Iteration, MetricName, Value, Unit, MetricType"],
    ],
)

doc.add_paragraph(
    "All endpoints return JSON. Input parameters are sanitized to prevent KQL injection."
)

# ── 8. Project Structure ────────────────────────────────────────────────────
doc.add_heading("8. Project Structure", level=1)

add_code_block(doc,
    "C:\\hobl-dashboard\\\n"
    "├── venv\\                      # Python virtual environment\n"
    "├── templates\\\n"
    "│   └── index.html             # Dashboard UI (HTML + CSS + JS)\n"
    "├── app.py                     # Flask backend with Kusto query logic\n"
    "├── config.py                  # Configuration (cluster, DB, table)\n"
    "├── requirements.txt           # Python dependencies\n"
    "└── DESIGN_DOCUMENT.md         # Markdown version of this document"
)

# ── 9. How to Run ────────────────────────────────────────────────────────────
doc.add_heading("9. How to Run", level=1)

doc.add_heading("Prerequisites", level=2)
doc.add_paragraph("Python 3.10+", style="List Bullet")
doc.add_paragraph(
    "VPN connection (to reach the Kusto cluster fungateprd.centralus)",
    style="List Bullet",
)
doc.add_paragraph(
    "Azure AD account with access to the FungatesDataStore Kusto database",
    style="List Bullet",
)

doc.add_heading("Steps", level=2)
add_code_block(doc,
    "cd C:\\hobl-dashboard\n"
    ".\\venv\\Scripts\\Activate.ps1\n"
    "python app.py"
)

doc.add_paragraph("1. The Flask server starts on http://127.0.0.1:5000", style="List Number")
doc.add_paragraph("2. Open that URL in your browser", style="List Number")
doc.add_paragraph(
    "3. A new tab opens for Azure AD login — pick your Microsoft account",
    style="List Number",
)
doc.add_paragraph(
    '4. After "Authentication complete", return to the dashboard tab',
    style="List Number",
)
doc.add_paragraph(
    "5. Use the cascading filters to explore metrics",
    style="List Number",
)

doc.add_heading("First-time Setup (if venv doesn't exist)", level=2)
add_code_block(doc,
    "cd C:\\hobl-dashboard\n"
    "python -m venv venv\n"
    ".\\venv\\Scripts\\Activate.ps1\n"
    "pip install -r requirements.txt\n"
    "python app.py"
)

# ── 10. Dependencies ────────────────────────────────────────────────────────
doc.add_heading("10. Dependencies", level=1)

add_table(doc,
    ["Package", "Version", "Purpose"],
    [
        ["flask", "3.1.1", "Web framework for the dashboard server"],
        ["azure-kusto-data", "4.6.1", "Kusto / Azure Data Explorer SDK"],
        ["azure-identity", "1.21.0", "Azure AD authentication (browser login)"],
    ],
)

# ── 11. Future Enhancements ─────────────────────────────────────────────────
doc.add_heading("11. Future Enhancements", level=1)

add_table(doc,
    ["Enhancement", "Description"],
    [
        ["CSV Export", "Add a button to download the displayed metrics as a CSV file"],
        ["Trend Charts", "Plot metric values across dates for the same device + scenario"],
        ["Comparison View", "Compare metrics between two different devices or dates side by side"],
        ["Team Deployment", "Host on Azure App Service with Azure AD web app auth for team-wide access"],
        ["Auto-Refresh", "Periodically poll for new data after scenario uploads"],
    ],
)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = r"C:\hobl-dashboard\HOBL_Dashboard_Design_Document.docx"
doc.save(output_path)
print(f"Word document saved to: {output_path}")
